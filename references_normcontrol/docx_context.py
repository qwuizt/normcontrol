from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from references_normcontrol.docx_tracked_editing import REFERENCE_END_RE, REFERENCE_TITLE_RE


@dataclass(frozen=True)
class DocxRunContext:
    run_index: int
    text: str
    bold: bool | None
    italic: bool | None
    underline: bool | None
    style_name: str | None


@dataclass(frozen=True)
class DocxParagraphContext:
    paragraph_index: int
    block_index: int
    text: str
    normalized_text: str
    text_hash: str
    style_name: str | None
    in_references: bool
    source: str
    table_index: int | None
    row_index: int | None
    cell_index: int | None
    runs: list[DocxRunContext]


@dataclass(frozen=True)
class DocxDocumentContext:
    path_docx: str
    paragraph_count: int
    table_count: int
    reference_start_paragraph_index: int | None
    reference_end_paragraph_index: int | None
    paragraphs: list[DocxParagraphContext]


BlockItem = Paragraph | Table


def normalize_docx_text(text: str) -> str:
    """Нормализовать DOCX-текст для поиска и устойчивых хэшей."""
    return re.sub(r'\s+', ' ', text.replace('\xa0', ' ')).strip().casefold()


def text_hash(text: str) -> str:
    return hashlib.sha1(normalize_docx_text(text).encode('utf-8')).hexdigest()


def build_docx_context(path_docx: Path) -> DocxDocumentContext:
    """
    Построить исследовательский индекс DOCX.

    Индекс сохраняет порядок абзацев, включая абзацы внутри таблиц, и
    присваивает ``paragraph_index`` в том же порядке, в котором абзацы идут в
    XML-документе. Это нужно для будущего сопоставления замечаний с местами
    правок.
    """
    document = Document(str(path_docx))
    paragraphs: list[DocxParagraphContext] = []
    table_count = 0
    paragraph_index = 0
    block_index = 0

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            paragraphs.append(
                build_paragraph_context(
                    block,
                    paragraph_index=paragraph_index,
                    block_index=block_index,
                    source='body',
                )
            )
            paragraph_index += 1
            block_index += 1
            continue

        table_index = table_count
        table_count += 1
        table_paragraphs = collect_table_paragraphs(
            block,
            table_index=table_index,
            start_paragraph_index=paragraph_index,
            block_index=block_index,
        )
        paragraphs.extend(table_paragraphs)
        paragraph_index += len(table_paragraphs)
        block_index += 1

    reference_start, reference_end = find_reference_range(paragraphs)
    reference_indexes = set()
    if reference_start is not None:
        reference_indexes = set(range(reference_start, reference_end or len(paragraphs)))
    paragraphs = [
        set_reference_flag(paragraph, paragraph.paragraph_index in reference_indexes) for paragraph in paragraphs
    ]

    return DocxDocumentContext(
        path_docx=str(path_docx),
        paragraph_count=len(paragraphs),
        table_count=table_count,
        reference_start_paragraph_index=reference_start,
        reference_end_paragraph_index=reference_end,
        paragraphs=paragraphs,
    )


def iter_block_items(parent: DocumentObject | _Cell) -> Iterator[BlockItem]:
    """
    Вернуть абзацы и таблицы верхнего уровня в порядке DOCX.

    Это стандартный низкоуровневый обход python-docx: публичное API
    ``document.paragraphs`` теряет позицию таблиц между абзацами.
    """
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
        parent_object = parent
    else:
        parent_element = parent._tc
        parent_object = parent

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent_object)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent_object)


def collect_table_paragraphs(
    table: Table,
    *,
    table_index: int,
    start_paragraph_index: int,
    block_index: int,
) -> list[DocxParagraphContext]:
    paragraphs: list[DocxParagraphContext] = []
    paragraph_index = start_paragraph_index

    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraphs.append(
                    build_paragraph_context(
                        paragraph,
                        paragraph_index=paragraph_index,
                        block_index=block_index,
                        source='table',
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=cell_index,
                    )
                )
                paragraph_index += 1

    return paragraphs


def build_paragraph_context(
    paragraph: Paragraph,
    *,
    paragraph_index: int,
    block_index: int,
    source: str,
    table_index: int | None = None,
    row_index: int | None = None,
    cell_index: int | None = None,
) -> DocxParagraphContext:
    text = paragraph.text
    normalized_text = normalize_docx_text(text)
    return DocxParagraphContext(
        paragraph_index=paragraph_index,
        block_index=block_index,
        text=text,
        normalized_text=normalized_text,
        text_hash=hashlib.sha1(normalized_text.encode('utf-8')).hexdigest(),
        style_name=paragraph.style.name if paragraph.style is not None else None,
        in_references=False,
        source=source,
        table_index=table_index,
        row_index=row_index,
        cell_index=cell_index,
        runs=[
            DocxRunContext(
                run_index=run_index,
                text=run.text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                style_name=run.style.name if run.style is not None else None,
            )
            for run_index, run in enumerate(paragraph.runs)
        ],
    )


def find_reference_range(paragraphs: list[DocxParagraphContext]) -> tuple[int | None, int | None]:
    start: int | None = None
    end: int | None = None

    for paragraph in paragraphs:
        if re.fullmatch(REFERENCE_TITLE_RE, paragraph.text.strip(), flags=re.IGNORECASE):
            start = paragraph.paragraph_index + 1
            break

    if start is None:
        return None, None

    for paragraph in paragraphs:
        if paragraph.paragraph_index < start:
            continue
        if re.match(REFERENCE_END_RE, paragraph.text.strip(), flags=re.IGNORECASE):
            end = paragraph.paragraph_index
            break

    return start, end or len(paragraphs)


def set_reference_flag(paragraph: DocxParagraphContext, in_references: bool) -> DocxParagraphContext:
    return DocxParagraphContext(
        paragraph_index=paragraph.paragraph_index,
        block_index=paragraph.block_index,
        text=paragraph.text,
        normalized_text=paragraph.normalized_text,
        text_hash=paragraph.text_hash,
        style_name=paragraph.style_name,
        in_references=in_references,
        source=paragraph.source,
        table_index=paragraph.table_index,
        row_index=paragraph.row_index,
        cell_index=paragraph.cell_index,
        runs=paragraph.runs,
    )


def save_docx_context(path_context: Path, context: DocxDocumentContext) -> None:
    payload = asdict(context)
    path_context.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
